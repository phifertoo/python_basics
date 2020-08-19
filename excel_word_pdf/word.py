# install pythoon-docx
# you cannot make changes to the existing document. You have to copy the text form the existing document,
# transfer the text to the new document with the changes

import docx

my_doc = docx.Document('demo.docx')

# the word document is separated by paragraphs. Each paragraph represents a "return"
first_para = my_doc.paragraphs[0].text
seoncd_para = my_doc.paragraphs[1].text

# a run is a difference in style (ie. bold, italics, underline)
first_run = my_doc.paragraphs[0].runs[0].text
second_run = my_doc.paragraphs[1].runs[0].text
third_run = my_doc.paragraphs[1].runs[1].text

# check if a run is bold
isBold = my_doc.paragraphs[1].runs[1].bold

# change text
added_text = my_doc.paragraphs[1].runs[1].text = 'my_new_text'
my_doc.save('demo.docx')

# create a new document
new_doc = docx.Document()
first_para1 = new_doc.add_paragraph('Hello this is my first paragraph')
second_para2 = new_doc.add_paragraph('Hello this is my second paragraph')
first_para1.add_run("This is a new run")
new_doc.save('new_doc.docx')

# get all the text from a word file
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))

